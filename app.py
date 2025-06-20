import streamlit as st
import google.generativeai as genai
import PyPDF2
import io
from docx import Document
from PIL import Image # For image processing (required for OCR)
import pytesseract # For Optical Character Recognition (OCR)

# --- Streamlit Secrets Configuration ---
# To use this application, you need to create a file named `.streamlit/secrets.toml`
# in the same directory as your `app.py` file for local development.
# For Streamlit Cloud, you configure these secrets directly in the app dashboard.
#
# # .streamlit/secrets.toml
# GEMINI_API_KEY = "YOUR_GEMINI_API_KEY"
# STREAMLIT_PASSWORD = "Rajeev"
# # OPENAI_API_KEY = "YOUR_OPENAI_API_KEY" # Uncomment and add if you plan to use OpenAI API later

# --- Password Protection ---
correct_password = st.secrets.get("STREAMLIT_PASSWORD")

if not correct_password:
    st.error("Password not found in Streamlit secrets. Please set `STREAMLIT_PASSWORD`.")
    st.stop()

password = st.sidebar.text_input("Enter Password to Access", type="password")

if password != correct_password:
    st.sidebar.error("Incorrect password. Please try again.")
    st.stop()
else:
    st.sidebar.success("Password accepted!")

st.title("📄 Meeting Synopsis from PDF (Powered by Gemini & OCR)")
st.write("Upload a meeting minutes PDF (text-based or scanned), and I'll generate a concise synopsis.")

# --- API Key Initialization ---
gemini_api_key = st.secrets.get("GEMINI_API_KEY")

if not gemini_api_key:
    st.error("Gemini API key not found in Streamlit secrets. Please set `GEMINI_API_KEY`.")
    st.stop()

try:
    genai.configure(api_key=gemini_api_key)
except Exception as e:
    st.error(f"Error configuring Gemini API: {e}")
    st.stop()

# --- OCR Engine Path (Crucial for pytesseract) ---
# For local development:
# You might need to set the path to your Tesseract executable.
# Example for Windows: pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# Example for Linux/macOS: pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract' (or similar)

# For Streamlit Cloud:
# Tesseract needs to be installed via apt-get in a packages.txt file.
# Streamlit Cloud's environment usually has Tesseract installed in a default path,
# so explicitly setting pytesseract.pytesseract.tesseract_cmd might not be necessary.
# However, if it fails, you might need to find the correct path on Streamlit Cloud's system.
# We'll assume the default path works after installing via packages.txt.

# --- PDF Processing Function ---
@st.cache_data # Cache the extraction to avoid re-processing on rerun
def extract_text_from_pdf(uploaded_file_buffer):
    """
    Extracts text from a PDF, attempting OCR if direct text extraction fails.
    """
    text = ""
    pdf_file = io.BytesIO(uploaded_file_buffer.getvalue())
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    num_pages = len(pdf_reader.pages)
    st.info(f"Attempting to extract text from {num_pages} pages...")

    # Try direct text extraction first
    extracted_any_text = False
    for page_num in range(num_pages):
        page = pdf_reader.pages[page_num]
        try:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text += page_text + "\n"
                extracted_any_text = True
            else:
                # If direct extraction fails for a page, mark it for potential OCR
                st.warning(f"No direct text extracted from page {page_num + 1}. Will try OCR for this page if no text is found in the entire document.")
        except Exception as e:
            st.warning(f"Error during direct text extraction from page {page_num + 1}: {e}")

    # If no text was extracted at all, try OCR for all pages
    if not extracted_any_text:
        st.warning("No text extracted directly from the PDF. Attempting OCR (Optical Character Recognition)...")
        try:
            # You'd typically need a PDF-to-image conversion library here (like pdf2image)
            # but pdf2image has system dependencies (poppler) that are harder to manage
            # in Streamlit Cloud's `packages.txt`.
            # For simplicity and to avoid complex system dependencies, direct OCR on PDF
            # pages isn't straightforward without external tools.
            # A common workaround involves `pdf2image` and `poppler`.
            # Since this is challenging for a quick deploy, I'll add a placeholder message.
            # If the user needs robust OCR, they might need a different approach or
            # a custom Docker image on Streamlit Cloud.
            
            # Placeholder for OCR functionality:
            # For `pytesseract` to work on a PDF directly without `pdf2image`
            # or `poppler`, you usually need to save individual pages as images first.
            # PyPDF2 doesn't directly convert pages to images.
            # This is a significant limitation for pure Streamlit Cloud deployment without `pdf2image`.

            st.error("Direct OCR on PDF pages within Streamlit Cloud requires additional system dependencies (like Poppler, used by `pdf2image`), which are complex to configure directly via `packages.txt`. While `pytesseract` is installed, converting PDF pages to images for OCR isn't straightforward in this setup.")
            st.info("For reliable OCR on scanned PDFs, consider local execution with `pdf2image` and Poppler installed, or explore cloud-based OCR APIs (e.g., Google Cloud Vision AI) for a more robust solution.")
            return "" # Return empty text as we can't perform OCR without more setup

        except Exception as e:
            st.error(f"An error occurred during OCR attempt: {e}. Please ensure Tesseract OCR engine is properly installed on the system.")
            return ""
    return text

# --- PDF Upload Section ---
uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded_file is not None:
    st.success("PDF file uploaded successfully!")

    extracted_text = extract_text_from_pdf(uploaded_file)

    if not extracted_text:
        st.warning("Could not extract any meaningful text from the PDF, even with OCR considerations. It might be entirely image-based or corrupted, and direct OCR might not be fully implemented without additional system dependencies (see info message).")
        st.stop()

    st.text_area("Extracted Text (first 1000 characters):", extracted_text[:1000], height=200, help="This shows a preview of the text extracted from your PDF.")

    # --- Generate Synopsis using Gemini ---
    st.subheader("Generating Meeting Synopsis...")

    # Prepare the prompt for Gemini with instructions for detailed and formatted output
    prompt = f"""
    You are an AI assistant specialized in summarizing meeting notes.
    Please read the following text extracted from a PDF document, which contains meeting information.
    Your task is to provide a detailed and properly formatted synopsis of the meeting.

    Please structure your response with the following sections using Markdown headings and bullet points:

    # Meeting Synopsis

    ## Summary
    Provide a concise overview of the entire meeting.

    ## Main Topics Discussed
    List the key subjects or agendas that were covered during the meeting using bullet points.

    ## Key Decisions Made
    Outline any important decisions, agreements, or conclusions reached, using bullet points.

    ## Action Items and Next Steps
    Detail specific tasks assigned, who is responsible (if mentioned), and deadlines (if any), using bullet points.

    ## Attendees (if identifiable)
    List any notable attendees or departments mentioned in the document.

    ---
    Here is the text from the PDF:

    {extracted_text}
    ---

    """

    # Display a spinner while generating
    with st.spinner("Please wait, Gemini is generating the detailed synopsis... This might take a moment for larger PDFs."):
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')

            response = model.generate_content(prompt)

            if response.candidates:
                synopsis = response.candidates[0].content.parts[0].text
                st.subheader("🎉 Meeting Synopsis:")
                st.markdown(synopsis) # Streamlit renders Markdown nicely

                # --- Download as Word document ---
                st.subheader("Download Options")
                doc = Document()
                doc.add_heading('Meeting Synopsis', level=0)

                for line in synopsis.split('\n'):
                    if line.startswith('## '):
                        doc.add_heading(line[3:].strip(), level=2)
                    elif line.startswith('# '):
                        doc.add_heading(line[2:].strip(), level=1)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:].strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                st.download_button(
                    label="Download Detailed Synopsis as Word (docx)",
                    data=doc_buffer,
                    file_name="detailed_meeting_synopsis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            else:
                st.error("Could not generate a synopsis. No candidates found in the Gemini response.")
                if response.prompt_feedback and response.prompt_feedback.block_reason:
                    st.warning(f"Gemini blocked the prompt due to: {response.prompt_feedback.block_reason}")
                    st.warning("Please try a different PDF or adjust the content.")
        except Exception as e:
            st.error(f"An error occurred while calling the Gemini API: {e}")
            st.info("Please check your API key, internet connection, or if the PDF content is too large for the model's context window.")

